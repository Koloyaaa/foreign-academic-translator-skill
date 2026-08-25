#!/usr/bin/env python3
"""
文本提取器：从 .pdf / .docx / .txt 中提取纯文本，供大模型处理。
质量门禁：若提取失败或内容为空，抛出明确异常并终止后续流程。
"""
import sys
import os
from pathlib import Path
from typing import Optional
import argparse

# 根据实际部署环境动态导入，若缺少库则给出明确提示
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

from utils import detect_encoding

def extract_from_pdf(file_path: Path) -> str:
    """从 PDF 提取文本，优先使用 pdfplumber（精度高），回退到 pypdf。"""
    if pdfplumber is not None:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                if text.strip():
                    return text.strip()
        except Exception as e:
            print(f"[WARN] pdfplumber 提取失败: {e}", file=sys.stderr)
    
    if pypdf is not None:
        try:
            reader = pypdf.PdfReader(file_path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if text.strip():
                return text.strip()
        except Exception as e:
            print(f"[WARN] pypdf 提取失败: {e}", file=sys.stderr)
    
    # 终极回退：尝试作为文本文件读取（某些 PDF 可能包含文本流）
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, "r", encoding=encoding, errors="ignore") as f:
            return f.read().strip()
    except Exception:
        pass
    
    raise RuntimeError(f"无法从 PDF 提取任何文本内容: {file_path}")

def extract_from_docx(file_path: Path) -> str:
    """从 Word 文档提取文本。"""
    if Document is None:
        raise ImportError("python-docx 未安装，无法解析 .docx 文件")
    
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text += "\n" + row_text
    if not full_text.strip():
        raise RuntimeError(f"提取的 .docx 内容为空: {file_path}")
    return full_text.strip()

def extract_from_txt(file_path: Path) -> str:
    """从纯文本文件提取内容。"""
    encoding = detect_encoding(file_path)
    with open(file_path, "r", encoding=encoding, errors="replace") as f:
        content = f.read().strip()
    if not content:
        raise RuntimeError(f"提取的 .txt 内容为空: {file_path}")
    return content

def extract(file_path: Path) -> str:
    """
    根据文件扩展名自动选择提取器。
    质量门禁：内容少于 10 个字符视为提取失败，抛出异常。
    """
    if not file_path.exists():
        raise FileNotFoundError(f"源文件不存在: {file_path}")
    
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        text = extract_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_from_docx(file_path)
    elif ext == ".txt":
        text = extract_from_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .pdf, .docx, .txt")
    
    # 质量门禁
    if len(text.strip()) < 10:
        raise RuntimeError(f"提取内容过短（长度 {len(text.strip())}），可能为无效文件或无法解析")
    
    return text

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="从学案文件中提取纯文本")
    parser.add_argument("input_path", type=str, help="输入文件路径 (.pdf/.docx/.txt)")
    parser.add_argument("--output", type=str, help="输出文本文件路径（可选）", default=None)
    args = parser.parse_args()
    
    try:
        result = extract(Path(args.input_path))
        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(result)
            print(f"文本已写入: {args.output}")
        else:
            print(result)
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)