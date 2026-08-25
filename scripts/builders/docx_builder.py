"""DOCX 构建器：将 Markdown 转换为 Word 文档（标题 / 段落 / 表格 / 页脚水印）。

支持 LaTeX 公式转 Word 原生公式（可编辑 oMath）：
  a) latex2mathml -> MathML -> XSLT(MML2OMML.XSL) -> OMML（首选）
  b) matplotlib mathtext -> 透明 PNG 图片（回退）
  c) 保持原样纯文本（最终回退）
依赖缺失时自动降级，build_docx 不抛异常中断。
"""

import re
import tempfile
from pathlib import Path

from .base import (
    WATERMARK,
    annotate_first_occurrences,
    protect_math,
    restore_math,
)

# ============ 公式转换后端（OMML -> PNG -> 纯文本 依次回退） ============

# 行内/行间公式正则：按 $$ → \[ → $ → \( 的优先级匹配（先匹配 $$...$$ 再匹配 \[...\]、
# $...$、\(...\)），避免 `$$` 被 `$` 误吞、`\\[`/`\\(` 被误判为转义；
# 公式内可能含反斜杠 \，故用非贪婪匹配并开启 DOTALL。
# 说明：`\[` 与 `\(` 在源文本中是两个字符（反斜杠 + 左方括号/左圆括号），此处用 raw string 并写成 `\\\[` / `\\\(` 使其按字面匹配。
_INLINE_MATH_RE = re.compile(
    r"\$\$(.+?)\$\$"        # $$...$$（行间）
    r"|\\\[(.+?)\\\]"       # \[...\]（行间，兼容兜底）
    r"|\$(.+?)\$"           # $...$（行内）
    r"|\\\((.+?)\\\)",      # \(...\)（行内，兼容兜底）
    re.DOTALL,
)
_DISPLAY_MATH_RE = re.compile(
    r"\$\$(.+?)\$\$"        # $$...$$（行间）
    r"|\\\[(.+?)\\\]",      # \[...\]（行间，兼容兜底）
    re.DOTALL,
)

# 后端状态：_omml_backend=(mathml_convert, xslt) 表示 OMML 可用；_image_backend 表示 matplotlib 可用
_omml_backend = None
_image_backend = False

# matplotlib 生成的临时图片，doc 保存后统一清理
_TEMP_IMAGE_PATHS = []


def _find_mml2omml_xsl():
    """依次尝试常见 Office 安装路径下的 MML2OMML.XSL，找不到返回 None（此后端不可用）。"""
    candidates = [
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\root\Office15\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office14\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office15\MML2OMML.XSL",
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None


def _init_formula_backends():
    """按可用性初始化公式后端：a) OMML（latex2mathml + lxml + MML2OMML.XSL）；b) matplotlib PNG。"""
    global _omml_backend, _image_backend
    _omml_backend = None
    _image_backend = False
    # 后端 a：latex2mathml 生成 MathML，lxml 加载 XSLT 转 OMML
    try:
        from latex2mathml.converter import convert as mathml_convert
        import lxml.etree as etree

        xsl_path = _find_mml2omml_xsl()
        if xsl_path:
            xslt = etree.XSLT(etree.parse(xsl_path))
            _omml_backend = (mathml_convert, xslt)
    except Exception:
        _omml_backend = None
    # 后端 b：matplotlib mathtext 渲染 PNG
    try:
        import matplotlib  # noqa: F401

        _image_backend = True
    except Exception:
        _image_backend = False


_init_formula_backends()


def _latex_to_omml_xml(latex: str, display: bool = False):
    """将 LaTeX 公式转换为 OMML（m:oMath）XML 字符串；失败返回 None。"""
    if _omml_backend is None:
        return None
    mathml_convert, xslt = _omml_backend
    try:
        import lxml.etree as etree

        # display='block' 让行间公式标记为 display 模式
        mml = mathml_convert(latex, display="block" if display else "inline")
        omml = xslt(etree.fromstring(mml.encode("utf-8")))
        return str(omml)
    except Exception:
        return None


def _insert_omml(paragraph, latex: str, display: bool = False) -> bool:
    """尝试以 Word 原生 oMath 形式插入公式；成功返回 True。"""
    xml = _latex_to_omml_xml(latex, display=display)
    if not xml:
        return False
    try:
        from docx.oxml import parse_xml

        omath = parse_xml(xml.encode("utf-8"))
        # 将 m:oMath 作为 w:p 的子元素按顺序追加，实现可编辑的原生公式
        paragraph._p.append(omath)
        return True
    except Exception:
        return False


def _render_math_png(latex: str, out_path: str) -> bool:
    """用 matplotlib mathtext 将公式渲染为透明 PNG 并裁剪空白；成功返回 True。"""
    if not _image_backend:
        return False
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        from matplotlib import mathtext
        import numpy as np
        import matplotlib.image as mpimg

        parser = mathtext.MathTextParser("agg")
        rgba, _depth = parser.to_rgba("$%s$" % latex, dpi=200)
        if rgba is None or rgba.size == 0:
            return False
        # 裁剪全透明边距，缩小图片尺寸
        alpha = rgba[:, :, 3]
        rows = np.any(alpha > 0, axis=1)
        cols = np.any(alpha > 0, axis=0)
        if rows.any() and cols.any():
            rgba = rgba[np.ix_(rows, cols)]
        mpimg.imsave(out_path, rgba, format="png")
        return True
    except Exception:
        return False


def _insert_math_image(paragraph, latex: str, display: bool = False) -> bool:
    """回退：用 matplotlib 渲染透明 PNG 并插入段落；成功返回 True。"""
    try:
        from docx.shared import Inches

        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close()
        if not _render_math_png(latex, tmp.name):
            try:
                Path(tmp.name).unlink(missing_ok=True)
            except Exception:
                pass
            return False
        _TEMP_IMAGE_PATHS.append(tmp.name)
        # 按像素宽度折算为英寸；行内公式略缩，行间公式保持原宽
        width_px = 640  # 兜底宽度
        try:
            from PIL import Image

            with Image.open(tmp.name) as im:
                width_px = im.size[0]
        except Exception:
            pass
        scale = 1.0 if display else 0.8
        run = paragraph.add_run()
        run.add_picture(tmp.name, width=Inches(width_px / 200.0 * scale))
        return True
    except Exception:
        return False


def _insert_math(paragraph, latex: str, display: bool = False) -> bool:
    """优先 OMML，其次 PNG 图片；均失败返回 False（调用方回退为纯文本）。"""
    if _insert_omml(paragraph, latex, display):
        return True
    return _insert_math_image(paragraph, latex, display)


# 不可见控制字符（\x00-\x08 \x0b \x0c \x0e-\x1f \x7f）与替换符（\ufffd）：写入 docx 前统一清理，避免把异常字符带进 Word 造成乱码
_INVALID_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\ufffd]")


def _sanitize_text(text: str) -> str:
    """移除不可见控制字符与替换符（替换为空字符串即可）。"""
    return _INVALID_CHARS_RE.sub("", text)


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False, highlight: bool = False):
    """添加一个 run，支持加粗 / 斜体 / 术语高亮（蓝青色加粗）；写入前清理异常控制字符。"""
    from docx.shared import RGBColor

    run = paragraph.add_run(_sanitize_text(text))
    if bold or highlight:
        run.bold = True
    if italic:
        run.italic = True
    if highlight:
        # 术语高亮：蓝青色加粗（与页面主色一致），不再使用黄色底纹
        try:
            run.font.color.rgb = RGBColor(0x2A, 0x7A, 0x9C)
        except Exception:
            pass
    return run


def _add_text_with_math(paragraph, text: str, bold: bool = False, italic: bool = False, highlight: bool = False):
    """将一段文本按公式切分：普通文本 -> run，公式 -> OMML/图片，保持原顺序。"""
    pos = 0
    for m in _INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], bold, italic, highlight)
        # $$...$$ 命中 group(1)，\[...\] 命中 group(2)，$...$ 命中 group(3)，\(...\) 命中 group(4)
        groups = m.groups()
        latex = next(g for g in groups if g is not None)
        display = m.group(1) is not None or m.group(2) is not None
        if not _insert_math(paragraph, latex, display):
            # 后端均不可用：保留原文（现状纯文本）
            _add_run(paragraph, m.group(0), bold, italic, highlight)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], bold, italic, highlight)


def _iter_children(node):
    """产出节点的子内容序列：str 文本 或 元素节点（对应 lxml.html 的 text/tail 语义）。"""
    if getattr(node, "text", None):
        yield node.text
    for child in node:
        yield child
        if getattr(child, "tail", None):
            yield child.tail


def _add_runs_with_math(paragraph, children, bold: bool = False, italic: bool = False, highlight: bool = False):
    """处理子节点序列：文本支持公式，strong/em/br/span 等元素递归处理；带 data-en 的术语 span 高亮加粗。"""
    for child in children:
        if isinstance(child, str):
            _add_text_with_math(paragraph, child, bold, italic, highlight)
            continue
        tag = getattr(child, "tag", None)
        if tag == "strong":
            _add_runs_with_math(paragraph, _iter_children(child), bold=True, italic=italic, highlight=highlight)
        elif tag == "em":
            _add_runs_with_math(paragraph, _iter_children(child), bold=bold, italic=True, highlight=highlight)
        elif tag == "br":
            paragraph.add_run("\n")
        else:
            # span 等其它内联元素：带 data-en 的术语 span 对其内容整体高亮加粗；其余递归取其子内容
            child_highlight = highlight or bool(getattr(child, "attrib", {}).get("data-en"))
            _add_runs_with_math(paragraph, _iter_children(child), bold=bold, italic=italic, highlight=child_highlight)


def _split_display_sections(el):
    """把段落节点按行间公式 $$...$$ 切成 (kind, payload) 序列；kind: text|math。"""
    sections = []
    text_parts = []

    def flush_text():
        if text_parts:
            sections.append(("text", list(text_parts)))
            del text_parts[:]

    for child in _iter_children(el):
        if isinstance(child, str):
            # _DISPLAY_MATH_RE 为双组交替 (组1)|(组2)，re.split 会对未命中组插入 None，
            # 先剔除以免奇偶错位把尾空串误判为 math 段（产生空公式 $$$$ 残留）
            parts = [p for p in _DISPLAY_MATH_RE.split(child) if p is not None]
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    if part:
                        text_parts.append(part)
                else:
                    flush_text()
                    sections.append(("math", part))
        else:
            text_parts.append(child)
    flush_text()
    return sections


def _add_h3_bar(para):
    """在 H3 段落最前插入琥珀色粗竖条装饰字符 run（▎ U+258E，加粗、16pt、颜色 F5B86E）。

    方案A：以字符代替原先的段落左边框（pBdr left），竖条更短（约一个字高）且更粗，
    视觉上更加精致；该装饰 run 需在每个 H3 段落生成时插入（样式级无法加字符）。
    """
    from docx.shared import Pt
    from docx.shared import RGBColor

    # 新建装饰 run 并移动到段落第一个 run 之前，确保竖条位于标题文本最前面
    bar_run = para.add_run("▎")
    if para.runs:
        first_run = para.runs[0]
        first_run._r.addprevious(bar_run._r)
    bar_run.bold = True
    try:
        bar_run.font.size = Pt(16)  # 字号略大于正文字号
        bar_run.font.color.rgb = RGBColor(0xF5, 0xB8, 0x6E)  # 琥珀色
    except Exception:
        pass


def _add_heading(doc, el):
    """h1/h2/h3 标题：按标签层级写入，样式由 _apply_styles 定义（边框/底纹/字体）。

    H3 额外在段落最前插入琥珀色竖条装饰字符（见 _add_h3_bar），
    替代原先的段落左边框（pBdr left），实现更短更粗的琥珀左条效果。
    """
    level = int(el.tag[1])
    para = doc.add_heading(_sanitize_text(el.text_content()), level=level)
    if level == 3:
        _add_h3_bar(para)


def _apply_reminder_style(para):
    """为段落添加提醒框样式：虚线边框 + 浅琥珀色背景 + 左右缩进 0.5cm。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = para._p.get_or_add_pPr()
    # 四边虚线边框（色 #F5B86E，大小 6pt）
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        edge_el = OxmlElement("w:%s" % edge)
        edge_el.set(qn("w:val"), "dashed")
        edge_el.set(qn("w:sz"), "6")
        edge_el.set(qn("w:space"), "4")
        edge_el.set(qn("w:color"), "F5B86E")
        p_bdr.append(edge_el)
    p_pr.append(p_bdr)
    # 浅琥珀色背景
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFAEB")
    p_pr.append(shd)
    # 左右缩进各 0.5cm（1cm ≈ 567 twips）
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "283")
    ind.set(qn("w:right"), "283")
    p_pr.append(ind)


def _add_paragraph(doc, el):
    """普通段落：支持公式。行内公式内联，行间公式（$$...$$ 独立成段）渲染为居中独立段落。"""
    full_text = el.text_content()
    # 检测提醒/教师要求等关键词，命中则添加虚线边框+浅琥珀色背景
    REMINDER_KEYWORDS = ["重要提醒", "注意", "教师要求", "老师提醒", "提醒", "请同学们", "要求"]
    is_reminder = any(kw in full_text for kw in REMINDER_KEYWORDS)

    if "$" not in full_text:
        para = doc.add_paragraph()
        _add_runs_with_math(para, _iter_children(el))
        if is_reminder:
            _apply_reminder_style(para)
        return
    # 段落含行间公式：拆为 文本段 / 居中公式段 / 文本段 ...
    for kind, payload in _split_display_sections(el):
        if kind == "text":
            para = doc.add_paragraph()
            _add_runs_with_math(para, payload)
        else:
            para = doc.add_paragraph()
            try:
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            if not _insert_math(para, payload, display=True):
                _add_run(para, "$$%s$$" % payload)
        if is_reminder:
            _apply_reminder_style(para)


def _style_header_cell(doc_cell):
    """表头单元格美化：柔和浅蓝底纹填充 DCE9F3（tcPr/shd），文字加粗深蓝 1F4E79，垂直居中。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL

    # 单元格垂直居中
    doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 单元格底纹（tcPr -> shd fill DCE9F3，柔和浅蓝）
    tc_pr = doc_cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "DCE9F3")
    tc_pr.append(shd)
    # 表头文字：加粗 + 深蓝色（柔和化，不再用纯白）
    for para in doc_cell.paragraphs:
        for run in para.runs:
            run.bold = True
            try:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            except Exception:
                pass


def _add_table(doc, el):
    """表格：解析 tr/td/th 生成 Word 表格；第一行视为表头（底纹+白字加粗），数据行偶数行交替底色，单元格内公式与 span 富文本一并处理。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # lxml.html 会在 table 内自动插入 tbody，故用 iter 取全部后代 tr
    rows = list(el.iter("tr"))
    if not rows:
        return
    ncols = max(
        len([c for c in tr if getattr(c, "tag", None) in ("td", "th")]) for tr in rows
    )
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    # 设置表格整体边框为细线 #D6DDE3（柔和浅灰）
    try:
        tbl_pr = table._tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_el = OxmlElement("w:%s" % edge)
            edge_el.set(qn("w:val"), "single")
            edge_el.set(qn("w:sz"), "4")
            edge_el.set(qn("w:space"), "0")
            edge_el.set(qn("w:color"), "D6DDE3")
            tbl_borders.append(edge_el)
        tbl_pr.append(tbl_borders)
    except Exception:
        pass
    for i, tr in enumerate(rows):
        cells = [c for c in tr if getattr(c, "tag", None) in ("td", "th")]
        is_header = i == 0  # 第一行视为表头
        for j, cell in enumerate(cells):
            if j < ncols:
                doc_cell = table.rows[i].cells[j]
                para = doc_cell.paragraphs[0]
                _add_runs_with_math(para, _iter_children(cell))
                if is_header:
                    _style_header_cell(doc_cell)
                elif i % 2 == 0:
                    # 偶数数据行（索引 2,4,6...）交替底色 #F7FAFC（更浅，柔和化）
                    try:
                        tc_pr = doc_cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "F7FAFC")
                        tc_pr.append(shd)
                    except Exception:
                        pass


# 标签 -> 处理函数 的映射，避免 if/elif 链
_TAG_HANDLERS = {
    "h1": _add_heading,
    "h2": _add_heading,
    "h3": _add_heading,
    "p": _add_paragraph,
    "table": _add_table,
}

# 需要处理的块级标签
_BLOCK_TAGS = frozenset(("h1", "h2", "h3", "p", "table"))


def _walk_blocks(el, fn):
    """前序遍历块级元素；命中目标标签即回调，不再深入其内部（表格/段落内部由各自处理器负责）。"""
    tag = getattr(el, "tag", None)
    if not isinstance(tag, str):
        return
    if tag in _BLOCK_TAGS:
        fn(el)
        return
    for child in el:
        _walk_blocks(child, fn)


def _parse_body(body_html: str):
    """解析 markdown 产出的 HTML 主体，返回顶层元素序列（兼容 lxml.html 的各种形态）。"""
    from lxml import html as lxml_html

    try:
        fragments = lxml_html.fragments_fromstring(body_html)
    except Exception:
        root = lxml_html.fromstring(body_html)
        fragments = list(root)
    return [f for f in fragments if not isinstance(f, str)]


def _apply_styles(doc):
    """为文档注入美观样式：正文/标题字体字号行距、标题边框/底纹/配色（中文字体同时设置 rFonts 的 w:eastAsia）。

    全部通过底层 XML（pPr/pBdr/shd、tcPr/shd、w:rPr）实现，保证 Word 兼容性。
    - H1：微软雅黑加粗 20pt #0B2A4A，底纹 #EBF5FB，底部+左侧粗边框
    - H2：微软雅黑加粗 16pt #1A4A6A，底纹 #EAF3F8，左侧边框
    - H3：微软雅黑加粗 14pt #2A5A7A，无底纹，左侧琥珀色装饰边框
    - 正文：等线 11pt #333333，1.3 倍行距，首行缩进 2 字符
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # pPr 子元素 schema 顺序，用于按正确位置插入，避免 Word 校验报错
    _PPR_ORDER = [
        "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
        "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
        "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap", "w:overflowPunct",
        "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
        "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
        "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
        "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr",
        "w:sectPr", "w:pPrChange",
    ]

    def _append_in_order(parent, el):
        """按 pPr 子元素 schema 顺序插入，保证 Word 兼容性。"""
        tag = el.tag
        if tag not in _PPR_ORDER:
            parent.append(el)
            return
        rank = _PPR_ORDER.index(tag)
        for child in list(parent):
            child_tag = child.tag
            if child_tag in _PPR_ORDER and _PPR_ORDER.index(child_tag) > rank:
                child.addprevious(el)
                return
        parent.append(el)

    def set_font_rpr(rpr, name, size_pt=None, bold=None, color=None):
        """在 w:rPr 上设置字体（ascii/eastAsia/hAnsi 同名）、字号、加粗、颜色。"""
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:eastAsia"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:cs"), name)
        if size_pt is not None:
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                rpr.append(sz)
            sz.set(qn("w:val"), str(int(size_pt * 2)))
            sz_cs = rpr.find(qn("w:szCs"))
            if sz_cs is None:
                sz_cs = OxmlElement("w:szCs")
                rpr.append(sz_cs)
            sz_cs.set(qn("w:val"), str(int(size_pt * 2)))
        if bold is not None:
            b = rpr.find(qn("w:b"))
            if b is None:
                b = OxmlElement("w:b")
                rpr.append(b)
            b.set(qn("w:val"), "1" if bold else "0")
        if color is not None:
            c = rpr.find(qn("w:color"))
            if c is None:
                c = OxmlElement("w:color")
                rpr.append(c)
            c.set(qn("w:val"), color)

    def set_spacing(pPr, before=None, after=None, line=None, first_line_chars=None):
        """在 pPr 上设置段前段后（单位：磅的 1/20）、行距（line 单位 + lineRule auto）与首行缩进（firstLineChars，200=2 字符）。"""
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            _append_in_order(pPr, spacing)
        if before is not None:
            spacing.set(qn("w:before"), str(before))
        if after is not None:
            spacing.set(qn("w:after"), str(after))
        if line is not None:
            spacing.set(qn("w:line"), str(line))
            spacing.set(qn("w:lineRule"), "auto")
        if first_line_chars is not None:
            spacing.set(qn("w:firstLineChars"), str(first_line_chars))

    def set_shd(pPr, fill):
        """在 pPr 上设置段落底纹填充色。"""
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        _append_in_order(pPr, shd)

    # ---------- 正文 Normal：等线 11pt，行距 1.3 倍，段后 6pt，正文颜色 #333333，首行缩进 2 字符 ----------
    normal = doc.styles["Normal"]
    set_font_rpr(normal._element.get_or_add_rPr(), "等线", size_pt=11, color="333333")
    set_spacing(normal._element.get_or_add_pPr(), before=0, after=120, line=300, first_line_chars=200)  # 段前 0、段后 6pt（120/20），行距 1.3 倍，首行缩进 2 字符

    # ---------- 标题 Heading 1/2/3 ----------
    # H1：微软雅黑加粗 20pt，色 #0B2A4A；底纹 #EBF5FB；段前 30pt 段后 10pt，行距 1.5 倍；底部+左侧边框 single sz=12 #2A7A9C
    h1 = doc.styles["Heading 1"]
    set_font_rpr(h1._element.get_or_add_rPr(), "微软雅黑", size_pt=20, bold=True, color="0B2A4A")
    set_spacing(h1._element.get_or_add_pPr(), before=600, after=200, line=360)  # 30pt/10pt，1.5 倍行距
    h1_pPr = h1._element.get_or_add_pPr()
    set_shd(h1_pPr, "EBF5FB")
    h1_pBdr = OxmlElement("w:pBdr")
    # 底部粗边框
    h1_bottom = OxmlElement("w:bottom")
    h1_bottom.set(qn("w:val"), "single")
    h1_bottom.set(qn("w:sz"), "12")
    h1_bottom.set(qn("w:space"), "4")
    h1_bottom.set(qn("w:color"), "2A7A9C")
    h1_pBdr.append(h1_bottom)
    # 左侧粗边框
    h1_left = OxmlElement("w:left")
    h1_left.set(qn("w:val"), "single")
    h1_left.set(qn("w:sz"), "12")
    h1_left.set(qn("w:space"), "4")
    h1_left.set(qn("w:color"), "2A7A9C")
    h1_pBdr.append(h1_left)
    _append_in_order(h1_pPr, h1_pBdr)

    # H2：微软雅黑加粗 16pt，色 #1A4A6A；底纹 #EAF3F8；段前 22pt 段后 10pt，行距 1.5 倍；左侧边框 single sz=8 #2A7A9C
    h2 = doc.styles["Heading 2"]
    set_font_rpr(h2._element.get_or_add_rPr(), "微软雅黑", size_pt=16, bold=True, color="1A4A6A")
    set_spacing(h2._element.get_or_add_pPr(), before=440, after=200, line=360)  # 22pt/10pt，1.5 倍行距
    h2_pPr = h2._element.get_or_add_pPr()
    set_shd(h2_pPr, "EAF3F8")
    h2_pBdr = OxmlElement("w:pBdr")
    # 左侧边框
    h2_left = OxmlElement("w:left")
    h2_left.set(qn("w:val"), "single")
    h2_left.set(qn("w:sz"), "8")
    h2_left.set(qn("w:space"), "4")
    h2_left.set(qn("w:color"), "2A7A9C")
    h2_pBdr.append(h2_left)
    _append_in_order(h2_pPr, h2_pBdr)

    # H3：微软雅黑加粗 14pt，色 #2A5A7A；段前 14pt 段后 8pt，行距 1.5 倍；无底纹、无段落边框
    # （琥珀色装饰改用段首竖条字符 ▎，见 _add_h3_bar；原左侧琥珀边框 pBdr left 已移除）
    h3 = doc.styles["Heading 3"]
    set_font_rpr(h3._element.get_or_add_rPr(), "微软雅黑", size_pt=14, bold=True, color="2A5A7A")
    set_spacing(h3._element.get_or_add_pPr(), before=280, after=160, line=360)  # 14pt/8pt，1.5 倍行距
    h3_pPr = h3._element.get_or_add_pPr()
    # 显式移除可能存在的遗留底纹与段落边框（确保 H3 无底纹、无 pBdr）
    for existing_shd in h3_pPr.findall(qn("w:shd")):
        h3_pPr.remove(existing_shd)
    for existing_bdr in h3_pPr.findall(qn("w:pBdr")):
        h3_pPr.remove(existing_bdr)


def build_docx(md_content: str, output_path: Path) -> Path:
    """将 Markdown 转换为 Word 文档。

    依赖：python-docx / markdown / lxml 为必需；latex2mathml 与 MML2OMML.XSL、matplotlib
    为公式转换的可选后端，缺失时自动降级，不会中断文档生成。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        import markdown
    except ImportError as e:
        raise ImportError("build_docx 缺少依赖（python-docx / markdown / lxml）: %s" % e)

    # 先保护公式再转 markdown，避免 _ / * 被转成 <em> 切碎公式；
    # 转换后还原，保证公式以完整 $$...$$ 形式进入 _split_display_sections / _add_text_with_math。
    protected_md, maths = protect_math(md_content)
    body_html = ""
    for exts in (
        ["tables", "fenced_code"],
        ["markdown.extensions.tables", "markdown.extensions.fenced_code"],
    ):
        try:
            body_html = markdown.markdown(protected_md, extensions=exts)
            break
        except (ImportError, KeyError):
            continue
    if not body_html:
        raise RuntimeError("build_docx 无法加载 markdown 扩展（tables / fenced_code）")

    body_html = annotate_first_occurrences(body_html)
    body_html = restore_math(body_html, maths)
    fragments = _parse_body(body_html)
    doc = Document()
    _apply_styles(doc)  # 注入美观样式：正文/标题字体字号行距、标题边框、配色

    for frag in fragments:
        _walk_blocks(frag, lambda el: _TAG_HANDLERS[el.tag](doc, el))

    # ---------- 页面设置与页码（页脚：页码域在前、换行后追加水印文本，合并到同一段避免页脚重复） ----------
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    # 页边距：上下 2cm，左右 2.5cm
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 页脚段落：水平居中
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入 PAGE 页码域（fldChar begin + instrText PAGE + fldChar end）
    run_fld = footer_para.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run_fld._r.append(fld_char1)
    run_fld._r.append(instr_text)
    run_fld._r.append(fld_char2)

    # 追加原水印文本 run（换行后），保留原有水印逻辑（字号 8pt）
    run = footer_para.add_run("\n欢迎使用 " + WATERMARK)
    run.font.size = Pt(8)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # 清理 matplotlib 生成的临时图片
    while _TEMP_IMAGE_PATHS:
        p = _TEMP_IMAGE_PATHS.pop()
        try:
            Path(p).unlink(missing_ok=True)
        except Exception:
            pass
    return output_path
