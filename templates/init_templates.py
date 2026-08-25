#!/usr/bin/env python3
"""
模板初始化脚本：由于 .docx 为二进制格式，无法通过纯文本直接创建，
此脚本用于生成空的占位模板文件，以保证项目结构完整。
用户可后续用自己设计的模板覆盖这些文件。
"""
import os
from pathlib import Path

# 尝试导入文档生成库，若未安装则给出提示
try:
    from docx import Document
except ImportError:
    Document = None

def init_word_template(output_path: Path) -> bool:
    """生成空白的 Word 模板（仅含一个标题占位符）。"""
    if Document is None:
        print("[WARN] python-docx 未安装，跳过 Word 模板生成。")
        return False
    
    doc = Document()
    doc.add_heading("学案译文模板", 0)
    doc.add_paragraph("此处将填入译文正文...")
    doc.save(str(output_path))
    print(f"[INFO] Word 模板已生成: {output_path}")
    return True

if __name__ == "__main__":
    # 确保在当前目录（templates/）下执行
    base_dir = Path(__file__).parent.absolute()
    
    word_path = base_dir / "word_template.docx"
    
    # 若文件已存在则不覆盖（防止用户自定义内容丢失）
    if word_path.exists():
        print(f"[SKIP] Word 模板已存在: {word_path}")
    else:
        init_word_template(word_path)
    
    print("\n[SUCCESS] 模板初始化完成。若需要自定义样式，请用设计工具编辑上述文件。")